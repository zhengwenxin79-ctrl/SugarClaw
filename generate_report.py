#!/usr/bin/env python3
"""生成 SugarClaw 专业分析报告 Word 文档"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# ─── 全局样式 ───────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.3
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Microsoft YaHei'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    if level == 1:
        hs.font.size = Pt(18)
        hs.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    elif level == 2:
        hs.font.size = Pt(14)
        hs.font.color.rgb = RGBColor(0x6C, 0x63, 0xFF)
    else:
        hs.font.size = Pt(12)
        hs.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9.5)
    # Data
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()


def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.2 + level * 0.8)
    p.paragraph_format.space_after = Pt(2)
    return p


def bold_bullet(title, desc):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.bold = True
    p.add_run(f' — {desc}')
    return p


# ═══════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('SugarClaw 专业分析报告')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x6C, 0x63, 0xFF)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('糖尿病智能管理平台 — 技术评估与改进方案')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run(f'报告日期：{datetime.date.today().strftime("%Y年%m月%d日")}')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)

doc.add_page_break()

# ═══════════════════════════════════════════
# 目录页
# ═══════════════════════════════════════════
doc.add_heading('目录', level=1)
toc_items = [
    '一、项目概述',
    '二、优点分析',
    '  2.1 多模型自适应卡尔曼滤波',
    '  2.2 对冲天平膳食风险评估',
    '  2.3 四层渐进式食物查询',
    '  2.4 瘦客户端 + 模块化架构',
    '  2.5 持久化与自适应校准',
    '三、不足分析与解决方案',
    '  3.1 安全与合规',
    '  3.2 临床准确性',
    '  3.3 可靠性与体验',
    '  3.4 运维与扩展',
    '四、综合评级',
    '五、路线图建议',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    if not item.startswith('  '):
        for run in p.runs:
            run.bold = True

doc.add_page_break()

# ═══════════════════════════════════════════
# 一、项目概述
# ═══════════════════════════════════════════
doc.add_heading('一、项目概述', level=1)

doc.add_paragraph(
    'SugarClaw 是一款面向糖尿病患者的智能血糖管理平台，采用 FastAPI 后端 + Flutter 跨平台前端的瘦客户端架构。'
    '系统集成了卡尔曼滤波血糖预测、AI 驱动食物 GI/GL 查询、对冲天平风险评估、CGM 模拟、PubMed 文献检索等核心模块。'
)

doc.add_heading('技术栈', level=3)
add_table(
    ['层级', '技术', '说明'],
    [
        ['后端', 'FastAPI + Python 3.9+', 'RESTful API + SSE 流式传输'],
        ['前端', 'Flutter 3.x (Dart)', '跨平台：Web / iOS / Android / Desktop'],
        ['数据库', 'SQLite (WAL 模式)', '5 张表：users / food_cache / cgm_readings / glucose_log / search_history'],
        ['AI', 'DeepSeek R1 / Chat', '对话引擎 + 食物数据兜底估算'],
        ['信号处理', 'NumPy + 自研 Kalman', 'KF / EKF / UKF 三种滤波器'],
        ['数据源', 'PubMed E-Utilities', '文献检索（stdlib，无额外依赖）'],
        ['状态管理', 'Provider (ChangeNotifier)', '6 个 Provider 管理全局状态'],
    ],
    col_widths=[3, 5, 8],
)

doc.add_heading('模块架构', level=3)
add_table(
    ['模块', '文件', '核心功能'],
    [
        ['卡尔曼引擎', 'kalman_engine.py', 'KF/EKF/UKF 自动选择、预测、CI、趋势、预警'],
        ['食物查询', 'query_food.py + foods_500.json', '精确匹配 → 拆词 → 向量搜索 → AI 兜底'],
        ['CGM 解析', 'ble_cgm_parser.py', 'BLE GATT 解析 + 24h 模拟数据生成'],
        ['PubMed', 'pubmed_researcher.py', 'esearch / esummary / efetch + 速率限制'],
        ['持久化', 'database.py', 'SQLite CRUD + 食物缓存 + 血糖日志'],
        ['对冲天平', 'api.py (scale 部分)', '风险量化 + 食物/运动/药物对冲方案生成'],
    ],
    col_widths=[3, 5, 8],
)

doc.add_page_break()

# ═══════════════════════════════════════════
# 二、优点分析
# ═══════════════════════════════════════════
doc.add_heading('二、优点分析', level=1)

# 2.1
doc.add_heading('2.1 多模型自适应卡尔曼滤波', level=2)
doc.add_paragraph(
    '系统实现了标准 KF、扩展 EKF、无迹 UKF 三种卡尔曼滤波器，'
    '并根据事件类型（进餐/胰岛素/运动/稳态）和输入参数自动选择最优滤波器。'
    '这在开源糖尿病管理工具中极为少见。'
)
bold_bullet('自动选择', '有胰岛素剂量时启用 EKF（胰岛素动力学非线性），有碳水数据时启用 UKF（碳水吸收建模），其余用标准 KF')
bold_bullet('UKF sigma-point', '使用 2n+1 个 sigma 点采样近似非线性碳水吸收函数，无需求解雅可比矩阵')
bold_bullet('95% 置信区间', '预测结果附带上下界，为临床决策提供不确定性量化')
bold_bullet('多层预警', 'CRITICAL / WARNING / PREDICTIVE 三级预警，含时间窗口预测')

# 2.2
doc.add_heading('2.2 对冲天平 — 独创膳食风险评估范式', level=2)
doc.add_paragraph(
    '将食物升糖风险量化为"重量"（0-100 分），通过低 GI 食物、运动、药物进行"对冲"，'
    '直觉化呈现了原本复杂的碳水代谢决策过程。'
)
bold_bullet('风险公式', 'base = (GL/50)×100，叠加纤维折扣（-15）、蛋白质折扣（-10）、脂肪减速（-5）、时间修正')
bold_bullet('时间感知', '6 个用餐时段动态调整：深夜加餐 +10、下午加餐 +5、正餐 +0')
bold_bullet('运动数据库', '90+ 运动条目含 MET 值（来源 Ainsworth 2011），支持模糊匹配和方言别名')
bold_bullet('地域亲和度', '优先推荐同地域低 GI 替代品（湖北用户优先看武汉本地食物）')

# 2.3
doc.add_heading('2.3 四层渐进式食物查询管线', level=2)
doc.add_paragraph('食物数据查询采用逐级回退策略，兼顾速度与覆盖率：')
add_table(
    ['层级', '方法', '延迟', '准确度'],
    [
        ['第 0 层', 'SQLite 缓存命中', '<1ms', '高（历史验证）'],
        ['第 1 层', '精确匹配 + 别名', '<5ms', '高（数据库条目）'],
        ['第 2 层', '拆词组合估算', '<10ms', '中高（加权合并）'],
        ['第 3 层', '向量语义搜索 (ChromaDB)', '<500ms', '中（距离阈值过滤）'],
        ['第 4 层', 'DeepSeek AI 估算', '2-5s', '中低（AI 生成，有幻觉风险）'],
    ],
    col_widths=[2.5, 5, 3, 5],
)
doc.add_paragraph('拆词算法能处理复合食物名：如"炸鸡咖喱饭"拆分为"炸鸡"+"咖喱"+"饭"，分别查询后加权合并营养数据。')

# 2.4
doc.add_heading('2.4 瘦客户端 + 模块化架构', level=2)
bold_bullet('零前端AI', 'Flutter 端不含任何 AI 逻辑，所有计算在后端完成，保证模型更新无需重新发布客户端')
bold_bullet('技能解耦', 'kalman_engine / ble_cgm_parser / pubmed_researcher / query_food 各自独立，可独立测试和替换')
bold_bullet('SSE 流式', 'Chat 对话和 CGM 回放使用 Server-Sent Events，实时推送增量数据')
bold_bullet('Agent 追踪', '每次分析记录各 Agent 的执行时间和结果，提供完整的决策透明度')

# 2.5
doc.add_heading('2.5 持久化与自适应校准', level=2)
bold_bullet('SQLite WAL', '写前日志模式保证崩溃恢复，外键约束保证数据完整性')
bold_bullet('ISF 自适应', 'EMA 公式 new_isf = 0.3×observed + 0.7×stored，每次校准自动更新用户胰岛素敏感因子')
bold_bullet('血糖日志', '支持任意时间点手动记录，不要求连续，适合真实用户碎片化记录习惯')
bold_bullet('用户画像注入', 'Chat 系统提示自动包含数据库中的用户档案，实现个性化对话')

doc.add_page_break()

# ═══════════════════════════════════════════
# 三、不足分析与解决方案
# ═══════════════════════════════════════════
doc.add_heading('三、不足分析与解决方案', level=1)

# ───── 3.1 ─────
doc.add_heading('3.1 安全与合规', level=2)

# 问题 1
doc.add_heading('问题 1：API 密钥泄露风险', level=3)
doc.add_paragraph(
    '当前状态：.env 文件含明文 DEEPSEEK_API_KEY，若被提交至 Git 仓库将导致密钥暴露。'
    '无密钥轮换机制，无过期处理。'
)
p = doc.add_paragraph()
run = p.add_run('解决方案：')
run.bold = True
run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
bullet('将 .env 加入 .gitignore，立即轮换已暴露的密钥')
bullet('使用操作系统环境变量或密钥管理服务（如 AWS Secrets Manager / HashiCorp Vault）')
bullet('添加密钥校验：启动时检查 API Key 格式和有效性，无效则拒绝启动并明确提示')
bullet('实现密钥轮换机制：支持热更新密钥而不重启服务')

# 问题 2
doc.add_heading('问题 2：零认证零授权', level=3)
doc.add_paragraph(
    '当前状态：无登录机制，无 Token 验证，CORS 全开（allow_origins=["*"]）。'
    '所有 API 端点裸露，任何客户端可访问全部健康数据。单用户硬编码 user_id=1。'
)
p = doc.add_paragraph()
run = p.add_run('解决方案：')
run.bold = True
run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
bullet('实现 JWT 认证：注册/登录 → 签发 access_token + refresh_token → 所有 API 校验 Bearer Token')
bullet('添加 FastAPI 依赖注入：创建 get_current_user() 依赖，所有端点自动校验身份')
bullet('数据库行级安全：所有查询加 WHERE user_id = current_user.id 条件')
bullet('CORS 限制为已知前端域名（如 https://sugarclaw.app）')
bullet('Flutter 端增加登录/注册页面，使用 SharedPreferences 存储 Token')

# 问题 3
doc.add_heading('问题 3：健康数据未加密', level=3)
doc.add_paragraph(
    '当前状态：SQLite 明文存储血糖、用药、ISF 等敏感健康数据。'
    '无数据保留策略，无 GDPR 删除接口。'
)
p = doc.add_paragraph()
run = p.add_run('解决方案：')
run.bold = True
run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
bullet('使用 SQLCipher 替代标准 SQLite，AES-256 加密数据库文件')
bullet('实现数据保留策略：CGM 读数默认保留 90 天，过期自动清理')
bullet('添加 DELETE /api/user/delete 端点：级联删除用户所有数据（GDPR 第 17 条）')
bullet('用户导出功能：GET /api/user/export 返回 JSON 格式的全部个人数据（GDPR 第 20 条）')
bullet('记录数据访问日志，便于合规审计')

doc.add_page_break()

# ───── 3.2 ─────
doc.add_heading('3.2 临床准确性', level=2)

# 问题 4
doc.add_heading('问题 4：卡尔曼滤波未经临床验证', level=3)
doc.add_paragraph(
    '当前状态：三种滤波器均基于理论模型实现，未与真实 CGM 设备（Dexcom G7 / FreeStyle Libre 3）数据对比验证。'
    '碳水吸收模型使用线性衰减（carb_remain × DT / t_decay），真实吸收为 S 型曲线。'
    'ISF 日内变异（黎明现象可使 ISF 波动 30-50%）未建模。传感器漂移（±10-15%）未纳入噪声模型。'
)
p = doc.add_paragraph()
run = p.add_run('解决方案：')
run.bold = True
run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
bullet('收集去标识化真实 CGM 数据（如公开数据集 OhioT1DM / D1NAMO），进行回测验证')
bullet('评估指标：MAE（目标 < 0.8 mmol/L）、MARD（目标 < 12%）、Clarke Error Grid 分析')
bullet('碳水吸收模型升级：替换为 Lehmann-Deutsch 双指数模型或 Hovorka 最小模型')
bullet('ISF 日内变异建模：引入时间加权 ISF = base_isf × circadian_factor(hour)，黎明 4-7 时 factor=1.3-1.5')
bullet('传感器噪声模型升级：从固定高斯噪声改为异方差模型 σ(g) = a + b×g（血糖越高噪声越大）')
bullet('发布验证报告，公开预测精度指标')

# 问题 5
doc.add_heading('问题 5：预警阈值硬编码且无个体化', level=3)
doc.add_paragraph(
    '当前状态：低血糖 3.9、紧急低 3.0、高血糖 10.0、紧急高 13.9 mmol/L 均为固定值。'
    '未考虑血糖下降速率、个体差异、场景敏感性。'
)
p = doc.add_paragraph()
run = p.add_run('解决方案：')
run.bold = True
run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
bullet('允许用户自定义阈值：在 Profile 页面增加"预警设置"，附带 ADA/EASD 推荐值作为默认')
bullet('引入变化率预警：dG/dt > 0.1 mmol/L/min 时触发"快速上升"警告，< -0.1 时触发"快速下降"')
bullet('场景敏感阈值：驾驶模式（低血糖阈值上调至 5.0）、睡眠模式（高血糖阈值放宽至 11.0）')
bullet('个体化学习：基于用户 30 天历史数据自动调整阈值百分位（如 P10 作为低血糖阈值）')

# 问题 6
doc.add_heading('问题 6：低血糖应急处理缺失', level=3)
doc.add_paragraph(
    '当前状态：触发低血糖警报后仅显示文本提示，无结构化处理方案，无低血糖频率追踪。'
)
p = doc.add_paragraph()
run = p.add_run('解决方案：')
run.bold = True
run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
bullet('实现 15/15 法则交互流程：当 BG < 4.0 → 弹出全屏警告 → "立即摄入 15g 速效糖（3-4 片葡萄糖片或 150ml 果汁）" → 15 分钟倒计时 → "复测血糖"')
bullet('低血糖事件记录：自动记录每次低血糖的时间、血糖值、处理方式、恢复时间')
bullet('频率追踪与红旗预警：若 7 天内 ≥ 2 次低血糖，提示"可能存在低血糖不自知风险，建议咨询内分泌科医师"')
bullet('紧急联系人：支持设置紧急联系人，严重低血糖（<3.0）时提供一键拨号')

# 问题 7
doc.add_heading('问题 7：食物数据库局限', level=3)
doc.add_paragraph(
    '当前状态：仅 500 种食物，中国菜系 10000+ 变体中覆盖率不足 5%。'
    'GI 值无来源标注、无更新日期。份量假设固定（如"米饭 150g"），未对接用户实际摄入量。'
)
p = doc.add_paragraph()
run = p.add_run('解决方案：')
run.bold = True
run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
bullet('扩充数据库：对接 USDA FoodData Central / 中国食物成分表（标准版），目标 3000+ 食物')
bullet('标注数据来源：每条食物记录添加 source（如"中国CDC 2020"）和 last_updated 字段')
bullet('GI 置信区间：从单值改为 [低, 中位, 高] 三值，如"米饭 GI = [64, 73, 92]"')
bullet('用户自定义份量：输入界面增加"实际摄入量(g)"字段，GL 动态计算 = GI × 实际碳水 / 100')
bullet('社区纠错：添加"GI 值似乎不准？"反馈按钮，后台审核后更新数据库')
bullet('烹饪方式修正：同一食材不同做法 GI 差异大（如煮土豆 GI=78 vs 炸薯条 GI=63），支持烹饪方式选择')

doc.add_page_break()

# ───── 3.3 ─────
doc.add_heading('3.3 可靠性与体验', level=2)

# 问题 8
doc.add_heading('问题 8：无离线能力', level=3)
doc.add_paragraph(
    '当前状态：所有分析依赖实时服务器调用。网络中断时应用完全不可用，无本地缓存。'
)
p = doc.add_paragraph()
run = p.add_run('解决方案：')
run.bold = True
run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
bullet('Flutter 端添加 sqflite 本地数据库，缓存最近的分析结果和食物 GI 数据')
bullet('离线模式降级：无网络时使用本地缓存的食物数据进行基础分析（不依赖 AI）')
bullet('自动同步：恢复网络后将离线期间记录的血糖日志批量上传')
bullet('添加网络状态指示器：AppBar 显示连接状态图标')

# 问题 9
doc.add_heading('问题 9：可观测性缺失', level=3)
doc.add_paragraph(
    '当前状态：使用 print 输出日志，无结构化日志、无指标采集、无告警。'
    '食物查询 4 层回退中命中哪层用户不知道。AI 估算数据与权威数据无标识区分。'
)
p = doc.add_paragraph()
run = p.add_run('解决方案：')
run.bold = True
run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
bullet('替换 print 为 Python logging 模块，配置 JSON 格式结构化日志')
bullet('添加 Prometheus 指标：请求延迟、错误率、食物查询各层命中率')
bullet('前端数据来源标识：食物卡片角标显示数据来源（"数据库" / "AI估算" / "向量匹配"）')
bullet('集成 Sentry 或类似服务捕获前后端异常')

# 问题 10
doc.add_heading('问题 10：前端错误处理粗糙', level=3)
doc.add_paragraph(
    '当前状态：所有 API 错误统一 throw Exception，不区分 404/500/超时。'
    '无重试逻辑，SSE 流无心跳检测。'
)
p = doc.add_paragraph()
run = p.add_run('解决方案：')
run.bold = True
run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
bullet('定义错误类型层级：NetworkError / ServerError / NotFoundError / TimeoutError')
bullet('添加指数退避重试：网络错误自动重试 3 次（1s → 2s → 4s）')
bullet('SSE 心跳检测：后端每 10 秒发送 heartbeat 事件，前端 15 秒未收到则重连')
bullet('用户友好错误信息：替换"Exception: ..."为"网络连接失败，请检查网络后重试"')

# 问题 11
doc.add_heading('问题 11：无障碍访问缺失', level=3)
doc.add_paragraph(
    '当前状态：颜色唯一编码（红=低血糖），色盲用户无法辨识。'
    '图表无 alt-text，无暗色模式，预警用 emoji 不利于屏幕阅读器。'
)
p = doc.add_paragraph()
run = p.add_run('解决方案：')
run.bold = True
run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
bullet('颜色 + 形状双编码：低血糖用红色 + 向下三角，高血糖用橙色 + 向上三角，正常用绿色 + 圆形')
bullet('为图表添加 Semantics 标签（Flutter Semantics widget），供 TalkBack/VoiceOver 朗读')
bullet('实现暗色模式：Material 3 ThemeData 添加 darkColorScheme')
bullet('替换 emoji 预警为 Material Icons + 文字描述')
bullet('WCAG 2.1 AA 对比度检查：确保所有文字与背景对比度 ≥ 4.5:1')

doc.add_page_break()

# ───── 3.4 ─────
doc.add_heading('3.4 运维与扩展', level=2)

# 问题 12
doc.add_heading('问题 12：部署与 CI/CD 空白', level=3)
doc.add_paragraph(
    '当前状态：无 Dockerfile，无 CI/CD 管线，无数据库迁移方案，无备份策略。'
)
p = doc.add_paragraph()
run = p.add_run('解决方案：')
run.bold = True
run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
bullet('编写 Dockerfile + docker-compose.yml：后端 + SQLite 卷挂载，一键启动')
bullet('GitHub Actions CI：push 触发 → lint → 单元测试 → 集成测试 → Flutter build web → 构建 Docker 镜像')
bullet('数据库迁移：引入 Alembic（SQLAlchemy）或自建版本号迁移（migration_001.sql, migration_002.sql ...）')
bullet('自动备份：cron 每日备份 sugarclaw.db 到对象存储（S3/MinIO），保留 30 天')
bullet('环境分离：dev / staging / prod 三套配置，通过环境变量切换')

# 问题 13
doc.add_heading('问题 13：性能瓶颈', level=3)
doc.add_paragraph(
    '当前状态：食物查询瀑布流最坏 10+ 秒。_ALL_FOODS 全局变量无线程安全锁。无 API 限流。'
)
p = doc.add_paragraph()
run = p.add_run('解决方案：')
run.bold = True
run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
bullet('食物查询并行化：第 1-3 层并行执行（asyncio.gather），取首个非空结果')
bullet('Redis 缓存层：热门食物查询结果缓存 24 小时（如部署到多实例时共享）')
bullet('_ALL_FOODS 线程安全：使用 threading.Lock 或启动时一次性加载后设为不可变')
bullet('API 限流：使用 slowapi 中间件，限制 10 req/s per IP')
bullet('UKF 优化：对 3 状态 sigma-point 预计算常量矩阵，减少重复运算')

doc.add_page_break()

# ═══════════════════════════════════════════
# 四、综合评级
# ═══════════════════════════════════════════
doc.add_heading('四、综合评级', level=1)

add_table(
    ['维度', '评级', '说明'],
    [
        ['算法创新', 'A', '三模 KF + 对冲天平 + 4 层食物查询，在开源同类中领先'],
        ['架构设计', 'B+', '模块化好，瘦客户端清晰，但缺离线 / 多用户 / 配置管理'],
        ['安全合规', 'D', '零认证、密钥泄露风险、数据未加密、无 GDPR 支持'],
        ['临床准确性', 'C', '算法有潜力但未验证，阈值不个体化，低血糖处理缺失'],
        ['前端体验', 'B', '功能完整、交互流畅，但无障碍缺失、错误处理弱'],
        ['运维就绪度', 'F', '无 CI/CD、无监控、无部署方案、无迁移策略'],
    ],
    col_widths=[3.5, 2, 10.5],
)

doc.add_heading('适用范围评估', level=3)
add_table(
    ['场景', '是否适用', '说明'],
    [
        ['学术研究 / 概念验证', '✓ 适用', '算法创新性和架构完整性足以支撑论文发表'],
        ['内部测试 / Demo', '✓ 适用', '功能完整，UI 美观，适合技术展示'],
        ['有医师监督的试用', '△ 有条件', '需补充知情同意流程和预警验证'],
        ['公开 App Store 发布', '✗ 不适用', '安全、合规、临床验证均未达标'],
        ['无监督临床使用', '✗ 不适用', '预测未验证，可能导致错误用药决策'],
    ],
    col_widths=[5, 3, 8],
)

doc.add_page_break()

# ═══════════════════════════════════════════
# 五、路线图建议
# ═══════════════════════════════════════════
doc.add_heading('五、路线图建议', level=1)

doc.add_heading('Phase 1：安全加固（1-2 周）', level=2)
bullet('密钥管理：迁移至环境变量，轮换已暴露密钥')
bullet('JWT 认证：实现注册/登录/Token 校验全流程')
bullet('数据加密：SQLCipher 替换标准 SQLite')
bullet('CORS 收紧 + API 限流')
bullet('知情同意弹窗：首次使用声明"非医疗建议"')

doc.add_heading('Phase 2：临床验证（2-4 周）', level=2)
bullet('收集公开 CGM 数据集，执行回测，发布 MAE / MARD / Clarke Grid 报告')
bullet('与内分泌科医师 Review 预警阈值')
bullet('实现低血糖 15/15 法则交互流程')
bullet('食物数据库扩充至 3000+，标注数据来源')
bullet('用户可自定义预警阈值和份量')

doc.add_heading('Phase 3：可靠性提升（2-3 周）', level=2)
bullet('Flutter 端 sqflite 离线缓存')
bullet('错误处理重构：类型化异常 + 重试 + 心跳')
bullet('结构化日志 + Prometheus 指标')
bullet('无障碍：颜色+形状双编码、暗色模式、屏幕阅读器支持')

doc.add_heading('Phase 4：运维就绪（1-2 周）', level=2)
bullet('Dockerfile + docker-compose')
bullet('GitHub Actions CI/CD 管线')
bullet('数据库迁移框架')
bullet('自动备份 + 监控告警')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 报告完 —')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)

# ─── 保存 ───────────────────────────────
output_path = '/Users/zwx/sugarclaw-app/SugarClaw分析报告.docx'
doc.save(output_path)
print(f'报告已保存至: {output_path}')
