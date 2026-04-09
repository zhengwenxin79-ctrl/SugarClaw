#!/usr/bin/env python3
"""生成 SugarClaw 技术学习指南 Word 文档"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

OUTPUT = os.path.join(os.path.expanduser("~"), "Downloads", "SugarClaw技术学习指南.docx")

BLUE = RGBColor(0x2E, 0x86, 0xC1)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x55, 0x55, 0x55)
GREEN = RGBColor(0x1E, 0x8E, 0x3E)
FONT_CN = "微软雅黑"
FONT_EN = "Calibri"
CODE_FONT = "Consolas"


def sf(run, size=10.5, bold=False, color=BLACK, font=FONT_EN):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)


def title(doc, text, size=22, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = align
    p.space_after = Pt(4)
    r = p.add_run(text)
    sf(r, size=size, bold=True, color=color)
    return p


def h1(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(18)
    p.space_after = Pt(6)
    r = p.add_run(text)
    sf(r, size=16, bold=True, color=BLUE)
    # bottom border
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single", qn("w:sz"): "8",
        qn("w:space"): "1", qn("w:color"): "2E86C1"})
    pBdr.append(bottom)
    pPr.append(pBdr)


def h2(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    p.space_after = Pt(4)
    r = p.add_run(text)
    sf(r, size=13, bold=True, color=BLACK)


def h3(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(2)
    r = p.add_run(text)
    sf(r, size=11, bold=True, color=GRAY)


def para(doc, text, size=10.5):
    p = doc.add_paragraph()
    p.space_before = Pt(2)
    p.space_after = Pt(2)
    r = p.add_run(text)
    sf(r, size=size)
    return p


def bullet(doc, text, size=10):
    p = doc.add_paragraph()
    p.space_before = Pt(1)
    p.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    r = p.add_run("• " + text)
    sf(r, size=size)


def code_block(doc, code, label=""):
    if label:
        p = doc.add_paragraph()
        p.space_before = Pt(4)
        p.space_after = Pt(1)
        r = p.add_run(label)
        sf(r, size=9, bold=True, color=GREEN)
    for line in code.strip().split("\n"):
        p = doc.add_paragraph()
        p.space_before = Pt(0)
        p.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)
        r = p.add_run(line)
        sf(r, size=8.5, color=GRAY, font=CODE_FONT)


def resource_box(doc, items):
    h3(doc, "推荐学习资源")
    for item in items:
        bullet(doc, item, size=9.5)


def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.2)
        s.right_margin = Cm(2.2)

    # ═══ 封面 ═══
    for _ in range(6):
        doc.add_paragraph()
    title(doc, "SugarClaw 技术学习指南", size=28)
    title(doc, "从零到一的全栈 AI 医学项目实战", size=16, color=GRAY)
    doc.add_paragraph()
    title(doc, "2026 年 3 月", size=12, color=GRAY)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("涵盖：FastAPI · Flutter · 卡尔曼滤波 · ChromaDB · LLM · PubMed · BLE · SQLite")
    sf(r, size=10, color=GRAY)
    doc.add_page_break()

    # ═══ 目录 ═══
    title(doc, "目录", size=18, color=BLACK, align=WD_ALIGN_PARAGRAPH.LEFT)
    toc = [
        "第一章  项目架构总览",
        "第二章  Python 后端开发（FastAPI）",
        "第三章  卡尔曼滤波器 — 血糖预测核心",
        "第四章  向量数据库与语义检索（ChromaDB + RAG）",
        "第五章  Flutter 前端开发",
        "第六章  SQLite 数据库设计",
        "第七章  大语言模型集成（DeepSeek API）",
        "第八章  PubMed 文献检索集成",
        "第九章  BLE 蓝牙协议与 CGM 数据",
        "第十章  部署与运维",
        "附录 A  推荐学习路径（12 周规划）",
        "附录 B  参考资源汇总",
    ]
    for t in toc:
        bullet(doc, t)
    doc.add_page_break()

    # ═══════════════════════════════════════
    # 第一章
    # ═══════════════════════════════════════
    h1(doc, "第一章  项目架构总览")

    h2(doc, "1.1 系统架构")
    para(doc, "SugarClaw 采用经典的前后端分离 + 技能模块层架构：")
    code_block(doc, """
┌─────────────────────────────────────────┐
│   Flutter 前端 (iOS / Android / Web)     │
│   MaterialDesign 3, Provider 状态管理     │
├─────────────── REST API ────────────────┤
│   FastAPI 后端 (Python 3.9+)             │
│   ├─ SQLite 持久化 (WAL 模式, 5 张表)     │
│   ├─ 权威指南知识库 (CDS/ADA/IDF/WHO)     │
│   └─ 多 Agent 协调器                     │
├─────────────── Skills 层 ───────────────┤
│   ├─ 卡尔曼滤波引擎 (KF / EKF / UKF)     │
│   ├─ 食物 GI/GL RAG (ChromaDB, 501食物)  │
│   ├─ PubMed 文献检索 (NCBI E-Utilities)  │
│   └─ BLE CGM 解析器 (GATT 0x1808)       │
└─────────────────────────────────────────┘
""", "架构图：")

    h2(doc, "1.2 技术栈全景")
    para(doc, "后端：Python 3.9+、FastAPI、Pydantic、SQLite3、OpenAI SDK（DeepSeek 兼容）")
    para(doc, "前端：Dart / Flutter 3.x、Provider、fl_chart、http、url_launcher")
    para(doc, "AI/ML：卡尔曼滤波（numpy）、ChromaDB 向量数据库、DeepSeek V3 LLM")
    para(doc, "数据：501 种中国食物 JSON、PubMed NCBI API、125 患者 CGM 校准数据集")
    para(doc, "部署：Cloudflare Tunnel 内网穿透、Flutter Web 静态文件挂载")

    h2(doc, "1.3 为什么选择这些技术？")
    bullet(doc, "FastAPI vs Flask/Django：自动生成 API 文档、原生 async、Pydantic 数据验证、SSE 支持好")
    bullet(doc, "Flutter vs React Native：单代码库同时编译 Web/iOS/Android，Dart 语言类型安全")
    bullet(doc, "SQLite vs PostgreSQL：零配置嵌入式，单文件部署，WAL 模式支持并发读，适合单机 MVP")
    bullet(doc, "ChromaDB vs Pinecone/Milvus：开源免费，纯 Python，适合小规模向量检索（<10K 条）")
    doc.add_page_break()

    # ═══════════════════════════════════════
    # 第二章
    # ═══════════════════════════════════════
    h1(doc, "第二章  Python 后端开发（FastAPI）")

    h2(doc, "2.1 FastAPI 框架入门")
    para(doc, "FastAPI 是一个高性能 Python Web 框架，基于 Starlette（ASGI）和 Pydantic（数据验证）。"
         "它的核心优势是自动生成 OpenAPI 文档、原生支持异步、以及通过 Pydantic 模型实现请求/响应的强类型验证。")

    h3(doc, "安装与第一个端点")
    code_block(doc, """
# 安装
pip install fastapi uvicorn

# app.py — 最小示例
from fastapi import FastAPI
from pydantic import BaseModel

app = FastAPI(title="我的第一个 API")

class HealthResponse(BaseModel):
    status: str
    message: str

@app.get("/api/health", response_model=HealthResponse)
async def health():
    return {"status": "ok", "message": "服务正常运行"}

# 启动: uvicorn app:app --reload --port 8080
# 访问: http://localhost:8080/docs 查看自动生成的 API 文档
""", "代码示例：")

    h2(doc, "2.2 Pydantic 请求模型")
    para(doc, "Pydantic 模型定义了 API 接收和返回的数据结构。FastAPI 自动验证请求数据，类型不匹配时返回 422 错误。")
    code_block(doc, """
from pydantic import BaseModel, Field
from typing import List, Optional

class AnalyzeRequest(BaseModel):
    readings: List[float] = Field(..., min_length=3, description="血糖读数序列")
    event: Optional[str] = Field(None, description="事件类型: meal/insulin/exercise")
    food: Optional[str] = None
    gi: float = 0
    gl: float = 0

@app.post("/api/analyze")
async def analyze(req: AnalyzeRequest):
    # req.readings 已经自动验证为 float 列表，且长度 >= 3
    return {"current_glucose": req.readings[-1]}
""", "代码示例：")

    h2(doc, "2.3 SSE 实时流")
    para(doc, "Server-Sent Events（SSE）是一种服务器向客户端单向推送数据的协议。"
         "相比 WebSocket，SSE 更简单（纯 HTTP）、自动重连、天然支持代理/CDN。"
         "SugarClaw 用 SSE 实现 Chat 流式回复和 CGM 数据推流。")
    code_block(doc, """
from fastapi.responses import StreamingResponse
import json, asyncio

@app.get("/api/stream")
async def stream_data():
    async def generate():
        for i in range(10):
            data = {"type": "data", "value": i}
            yield f"data: {json.dumps(data)}\\n\\n"
            await asyncio.sleep(0.5)
        yield f"data: {json.dumps({'type': 'done'})}\\n\\n"

    return StreamingResponse(generate(), media_type="text/event-stream")
""", "SSE 服务端示例：")

    h2(doc, "2.4 CORS 与静态文件")
    para(doc, "跨域资源共享（CORS）允许前端从不同域名访问 API。Flutter Web 构建后的产物可直接挂载为静态文件。")
    code_block(doc, """
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles

app.add_middleware(CORSMiddleware, allow_origins=["*"],
                   allow_methods=["*"], allow_headers=["*"])

# Flutter build web 的产物挂载到根路径
app.mount("/", StaticFiles(directory="../frontend/build/web", html=True))
""", "代码示例：")

    resource_box(doc, [
        "FastAPI 官方文档: https://fastapi.tiangolo.com/zh/",
        "Pydantic V2 文档: https://docs.pydantic.dev/",
        "视频教程: freeCodeCamp - FastAPI Course (YouTube, 免费)",
    ])
    doc.add_page_break()

    # ═══════════════════════════════════════
    # 第三章
    # ═══════════════════════════════════════
    h1(doc, "第三章  卡尔曼滤波器 — 血糖预测核心")

    h2(doc, "3.1 什么是卡尔曼滤波？")
    para(doc, "卡尔曼滤波是一种递归算法，用于从含噪声的观测数据中估计系统的真实状态。"
         "直觉理解：你每天称体重，体重秤有 ±0.5kg 误差。卡尔曼滤波结合'昨天的体重'（预测）"
         "和'今天秤上的读数'（观测），给出一个比单独看任何一个都更准确的估计。")

    h2(doc, "3.2 标准 KF 的 5 个方程")
    para(doc, "卡尔曼滤波分为两步：预测（Predict）和更新（Update）。")

    h3(doc, "预测步骤")
    code_block(doc, """
x_pred = F @ x          # 状态预测：x(k|k-1) = F * x(k-1|k-1)
P_pred = F @ P @ F.T + Q  # 协方差预测：P(k|k-1) = F * P * F^T + Q

其中：
  x = 状态向量 [血糖值, 变化率]
  F = 状态转移矩阵（描述血糖如何随时间变化）
  P = 估计误差协方差（我们对估计值的不确定度）
  Q = 过程噪声（系统本身的随机扰动，如激素波动）
""", "公式含义：")

    h3(doc, "更新步骤")
    code_block(doc, """
K = P_pred @ H.T @ inv(H @ P_pred @ H.T + R)  # 卡尔曼增益
x = x_pred + K @ (z - H @ x_pred)              # 状态更新
P = (I - K @ H) @ P_pred                        # 协方差更新

其中：
  z = 观测值（CGM 读数）
  H = 观测矩阵（将状态映射到观测空间）
  R = 观测噪声（CGM 传感器误差）
  K = 卡尔曼增益（决定信任预测还是观测的权重）

关键直觉：
  R 大（传感器噪声大）→ K 小 → 更信任预测
  Q 大（系统不稳定）  → K 大 → 更信任观测
""", "公式含义：")

    h2(doc, "3.3 为什么需要 EKF 和 UKF？")
    para(doc, "标准 KF 假设系统是线性的。但血糖动力学是非线性的：")
    bullet(doc, "胰岛素衰减：指数衰减模型 glucose -= dose * ISF * exp(-dt/tau)，tau=77min")
    bullet(doc, "碳水吸收：先升后降的峰值模型 glucose += effect * (t/t_peak) * exp(1 - t/t_peak)")
    bullet(doc, "运动效应：强度相关的非线性葡萄糖消耗")

    h3(doc, "EKF（扩展卡尔曼）")
    para(doc, "EKF 用雅可比矩阵（偏导数）对非线性函数做局部线性化。"
         "适用于胰岛素注射和运动场景——这些场景的非线性是平滑的。")
    code_block(doc, """
# EKF 的非线性状态转移（胰岛素场景）
def f(x):
    glucose, rate, insulin_active = x
    # 胰岛素指数衰减
    decay = math.exp(-DT / tau)
    new_insulin = insulin_active * decay
    consumed = insulin_active * (1 - decay)
    # 血糖下降 = 消耗的胰岛素 * 敏感因子
    glucose_change = rate * DT - consumed * ISF
    return [glucose + glucose_change, rate * 0.95, new_insulin]
""", "代码示例：")

    h3(doc, "UKF（无迹卡尔曼）")
    para(doc, "UKF 不做线性化，而是选择一组'Sigma 点'通过非线性函数，"
         "然后用加权平均恢复统计特性。适用于餐后血糖预测——碳水吸收的非线性更强。")
    code_block(doc, """
# UKF Sigma 点生成
def sigma_points(x, P, alpha=1e-3, beta=2, kappa=0):
    n = len(x)
    lam = alpha**2 * (n + kappa) - n
    sigmas = [x]
    sqrt_P = np.linalg.cholesky((n + lam) * P)
    for i in range(n):
        sigmas.append(x + sqrt_P[i])
        sigmas.append(x - sqrt_P[i])
    return sigmas  # 共 2n+1 个点
""", "代码示例：")

    h2(doc, "3.4 参数校准的重要性")
    para(doc, "SugarClaw 基于 125 名患者、128,157 条 CGM 数据校准了关键参数：")
    bullet(doc, "观测噪声 R：默认 0.25 → 校准后 5.042（+1917%）— CGM 传感器比想象中噪声大得多")
    bullet(doc, "过程噪声 Q：默认 0.01 → 校准后 0.004（-57%）— 人体血糖变化比想象中平稳")
    bullet(doc, "胰岛素 tau：默认 55min → 校准后 77min（+40%）— 胰岛素作用比教科书慢")
    bullet(doc, "ISF：默认 2.5 → 校准后 0.73（-71%）— 真实胰岛素敏感性远低于理论值")
    para(doc, "最终 Clarke Error Grid A 区准确率：95.68%（RMSE 12.35 mg/dL）")

    h2(doc, "3.5 简化版 KF 实现（20 行 Python）")
    code_block(doc, """
import numpy as np

def kalman_filter(readings, Q=0.004, R=5.0):
    \"\"\"最简卡尔曼滤波器：输入血糖读数序列，返回滤波后的序列。\"\"\"
    dt = 5.0  # 采样间隔 5 分钟
    # 状态: [血糖, 变化率]
    x = np.array([readings[0], 0.0])
    P = np.diag([R, 0.1])
    F = np.array([[1, dt], [0, 1]])       # 状态转移
    H = np.array([[1, 0]])                 # 观测矩阵
    Q_mat = np.diag([Q, Q * 0.1])
    R_mat = np.array([[R]])
    I = np.eye(2)
    filtered = []
    for z in readings:
        # Predict
        x = F @ x
        P = F @ P @ F.T + Q_mat
        # Update
        S = H @ P @ H.T + R_mat
        K = P @ H.T @ np.linalg.inv(S)
        x = x + (K @ np.array([[z - H @ x]])).flatten()
        P = (I - K @ H) @ P
        filtered.append(float(x[0]))
    return filtered

# 使用示例
readings = [6.5, 7.0, 8.2, 9.8, 11.5, 13.2, 14.8]
result = kalman_filter(readings)
print(result)  # 平滑后的血糖序列
""", "完整可运行代码：")

    resource_box(doc, [
        "《Kalman and Bayesian Filters in Python》— 免费电子书 (github.com/rlabbe/Kalman-and-Bayesian-Filters-in-Python)",
        "3Blue1Brown — Bayes 定理视频 (建立直觉)",
        "论文: Turksoy et al. 'Multivariable Adaptive Identification and Control for Artificial Pancreas Systems'",
        "Clarke Error Grid 介绍: Clarke WL et al., Diabetes Care, 1987",
    ])
    doc.add_page_break()

    # ═══════════════════════════════════════
    # 第四章
    # ═══════════════════════════════════════
    h1(doc, "第四章  向量数据库与语义检索（ChromaDB + RAG）")

    h2(doc, "4.1 什么是向量嵌入（Embedding）？")
    para(doc, "向量嵌入是将文本/图片等非结构化数据转换为固定维度的数字向量的过程。"
         "相似含义的文本会被映射到向量空间中相近的位置。例如：")
    bullet(doc, '"热干面" → [0.23, -0.15, 0.87, ...] (384维向量)')
    bullet(doc, '"武汉早餐面" → [0.21, -0.14, 0.85, ...] (相近！)')
    bullet(doc, '"红烧排骨" → [-0.45, 0.32, 0.12, ...] (较远)')

    h2(doc, "4.2 ChromaDB 入门")
    para(doc, "ChromaDB 是一个轻量级开源向量数据库，纯 Python 实现，支持嵌入式部署。")
    code_block(doc, """
import chromadb

# 创建/连接持久化数据库
client = chromadb.PersistentClient(path="./chromadb")

# 创建集合（自带默认嵌入模型）
collection = client.get_or_create_collection("food_db")

# 添加文档
collection.add(
    documents=["热干面 武汉特色面食 GI=82", "白米饭 主食 GI=83"],
    metadatas=[{"gi": 82, "region": "湖北"}, {"gi": 83, "region": "全国"}],
    ids=["food_001", "food_002"]
)

# 语义查询
results = collection.query(
    query_texts=["过早吃了碗面"],  # 方言也能匹配！
    n_results=3
)
print(results["documents"])  # → [["热干面 武汉特色面食 GI=82", ...]]
""", "代码示例：")

    h2(doc, "4.3 RAG（检索增强生成）架构")
    para(doc, "RAG = Retrieval-Augmented Generation，即先检索相关知识，再交给 LLM 生成回答。"
         "这解决了 LLM 知识过时、幻觉的问题。SugarClaw 的 RAG 流程：")
    bullet(doc, "用户输入 '热干面' → ChromaDB 语义搜索 → 找到 GI=82, GL=65 等数据")
    bullet(doc, "将检索到的食物数据注入 LLM 的 prompt → LLM 生成个性化建议")
    bullet(doc, "分层检索策略：精确匹配(优先级0) → 别名匹配(1) → 子串匹配(2) → 语义搜索(3)")

    h2(doc, "4.4 SugarClaw 的分层检索策略")
    para(doc, "为什么不只用向量搜索？因为精确匹配更快更准，语义搜索是兜底。")
    code_block(doc, """
def query_food(query):
    # Layer 0: SQLite 缓存（最快）
    cached = database.get_cached_food(query)
    if cached: return cached

    # Layer 1: 精确名称匹配（食物名/别名）
    exact = exact_match(query, all_foods)
    if exact: return exact

    # Layer 2: 向量语义搜索（ChromaDB）
    results = collection.query(query_texts=[query], n_results=5)
    if results and results["distances"][0][0] < 1.2:
        return parse_results(results)

    # Layer 3: DeepSeek AI 估算（最后手段）
    return ai_estimate(query)
""", "分层检索伪代码：")

    resource_box(doc, [
        "ChromaDB 官方文档: https://docs.trychroma.com/",
        "LangChain RAG 教程: https://python.langchain.com/docs/tutorials/rag/",
        "Sentence Transformers (嵌入模型): https://www.sbert.net/",
    ])
    doc.add_page_break()

    # ═══════════════════════════════════════
    # 第五章
    # ═══════════════════════════════════════
    h1(doc, "第五章  Flutter 前端开发")

    h2(doc, "5.1 Dart 语言速成")
    para(doc, "Dart 是 Flutter 的编程语言，语法类似 Java/JavaScript，支持空安全和异步。")
    code_block(doc, """
// 变量与类型
String name = '热干面';
int gi = 82;
double? isf = null;  // 可空类型
List<double> readings = [6.5, 7.0, 8.2];

// 类
class FoodItem {
  final String name;
  final int gi;
  FoodItem({required this.name, required this.gi});

  factory FoodItem.fromJson(Map<String, dynamic> json) {
    return FoodItem(name: json['name'], gi: json['gi']);
  }
}

// 异步
Future<String> fetchData() async {
  final response = await http.get(Uri.parse('http://api/health'));
  return response.body;
}
""", "Dart 基础语法：")

    h2(doc, "5.2 Provider 状态管理")
    para(doc, "Provider 是 Flutter 官方推荐的轻量级状态管理方案。"
         "核心思路：ChangeNotifier 持有状态 → notifyListeners() 通知 UI 刷新 → Consumer 监听变化。")
    code_block(doc, """
// 1. 定义状态类
class ScaleState extends ChangeNotifier {
  double _riskWeight = 0;
  double get riskWeight => _riskWeight;

  Future<void> calculateRisk(String food) async {
    final result = await ApiService().calculateRisk(food);
    _riskWeight = result.riskWeight;
    notifyListeners();  // 通知 UI 刷新
  }
}

// 2. 在 main.dart 注册
MultiProvider(
  providers: [
    ChangeNotifierProvider(create: (_) => ScaleState()),
    ChangeNotifierProvider(create: (_) => ChatState()),
  ],
  child: MyApp(),
)

// 3. 在 Widget 中消费
Consumer<ScaleState>(
  builder: (context, state, _) => Text('风险: ${state.riskWeight}'),
)
""", "代码示例：")

    h2(doc, "5.3 网络请求与 JSON 解析")
    code_block(doc, """
import 'dart:convert';
import 'package:http/http.dart' as http;

Future<Map<String, dynamic>> postAnalyze(List<double> readings) async {
  final response = await http.post(
    Uri.parse('http://localhost:8080/api/analyze'),
    headers: {'Content-Type': 'application/json'},
    body: json.encode({'readings': readings}),
  );
  if (response.statusCode == 200) {
    return json.decode(response.body);
  }
  throw Exception('请求失败: ${response.statusCode}');
}
""", "HTTP POST 示例：")

    h2(doc, "5.4 SSE 流消费（Chat 实时回复）")
    code_block(doc, """
Future<void> streamChat(String message, Function(String) onToken) async {
  final request = http.Request('POST', Uri.parse('/api/chat'));
  request.headers['Content-Type'] = 'application/json';
  request.body = json.encode({'messages': [{'role': 'user', 'content': message}]});

  final response = await http.Client().send(request);
  final stream = response.stream
      .transform(utf8.decoder)
      .transform(const LineSplitter());

  await for (final line in stream) {
    if (!line.startsWith('data: ')) continue;
    final data = json.decode(line.substring(6));
    if (data['type'] == 'content') onToken(data['content']);
    if (data['type'] == 'done') break;
  }
}
""", "SSE 客户端消费：")

    h2(doc, "5.5 Flutter Web 部署")
    code_block(doc, """
# 1. 构建 Web 产物
flutter build web --release

# 2. 产物在 frontend/build/web/ 目录
# 3. FastAPI 挂载为静态文件（已在 api.py 中配置）
# 4. 访问 http://localhost:8080 即可
""", "部署步骤：")

    resource_box(doc, [
        "Flutter 官方 Codelabs: https://docs.flutter.dev/codelabs",
        "Dart 语言之旅: https://dart.dev/language",
        "Provider 文档: https://pub.dev/packages/provider",
        "fl_chart (图表库): https://pub.dev/packages/fl_chart",
    ])
    doc.add_page_break()

    # ═══════════════════════════════════════
    # 第六章
    # ═══════════════════════════════════════
    h1(doc, "第六章  SQLite 数据库设计")

    h2(doc, "6.1 为什么选择 SQLite？")
    bullet(doc, "零配置：不需要安装数据库服务器，Python 自带 sqlite3 模块")
    bullet(doc, "单文件：整个数据库就是一个 .db 文件，备份就是复制文件")
    bullet(doc, "性能：对于读多写少的场景，SQLite 比 PostgreSQL 更快（无网络开销）")
    bullet(doc, "WAL 模式：Write-Ahead Logging 允许读写并发，不会互相阻塞")

    h2(doc, "6.2 SugarClaw 的 5 张表")
    code_block(doc, """
-- 用户档案
CREATE TABLE users (
    id INTEGER PRIMARY KEY,
    name TEXT, age INTEGER, weight REAL, height REAL,
    diabetes_type TEXT,          -- '1型' / '2型' / '妊娠期'
    medications TEXT DEFAULT '[]',  -- JSON 数组
    isf REAL DEFAULT 0,          -- 胰岛素敏感系数
    icr REAL DEFAULT 0           -- 碳水比
);

-- 食物缓存（避免重复调用 AI）
CREATE TABLE food_cache (
    food_name TEXT UNIQUE NOT NULL,
    gi_value REAL, gl_per_serving REAL,
    carb_g REAL, protein_g REAL, fat_g REAL, fiber_g REAL,
    regional_tag TEXT DEFAULT '全国'
);

-- CGM 读数（时间序列）
CREATE TABLE cgm_readings (
    session_id TEXT NOT NULL,
    timestamp TEXT NOT NULL,
    glucose_mmol REAL NOT NULL,
    event TEXT DEFAULT ''
);
CREATE INDEX idx_cgm_session ON cgm_readings(session_id);

-- 血糖日志（手动记录）
CREATE TABLE glucose_log (
    timestamp TEXT NOT NULL,
    glucose_mmol REAL NOT NULL,
    note TEXT DEFAULT ''
);

-- PubMed 搜索历史
CREATE TABLE search_history (
    query TEXT NOT NULL,
    results_json TEXT DEFAULT '[]'
);
""", "建表语句：")

    h2(doc, "6.3 Python sqlite3 CRUD 示例")
    code_block(doc, """
import sqlite3

def _conn():
    conn = sqlite3.connect("sugarclaw.db")
    conn.row_factory = sqlite3.Row    # 行转字典
    conn.execute("PRAGMA journal_mode=WAL")  # WAL 模式
    return conn

# 查询
def get_user(user_id=1):
    conn = _conn()
    row = conn.execute("SELECT * FROM users WHERE id = ?", (user_id,)).fetchone()
    conn.close()
    return dict(row) if row else None

# 插入（参数化防 SQL 注入）
def save_glucose(timestamp, glucose, note=""):
    conn = _conn()
    conn.execute(
        "INSERT INTO glucose_log (timestamp, glucose_mmol, note) VALUES (?, ?, ?)",
        (timestamp, glucose, note)  # 永远用 ? 占位符，不要拼接字符串！
    )
    conn.commit()
    conn.close()
""", "代码示例：")

    resource_box(doc, [
        "SQLite 官方文档: https://www.sqlite.org/docs.html",
        "DB Browser for SQLite (可视化工具): https://sqlitebrowser.org/",
        "Python sqlite3 模块: https://docs.python.org/3/library/sqlite3.html",
    ])
    doc.add_page_break()

    # ═══════════════════════════════════════
    # 第七章
    # ═══════════════════════════════════════
    h1(doc, "第七章  大语言模型集成（DeepSeek API）")

    h2(doc, "7.1 OpenAI 兼容接口")
    para(doc, "DeepSeek、Moonshot、GLM 等国产大模型都兼容 OpenAI 的 API 格式。"
         "这意味着你学会一套接口，可以无缝切换不同模型。")
    code_block(doc, """
from openai import OpenAI

client = OpenAI(
    api_key="sk-你的密钥",
    base_url="https://api.deepseek.com"  # 换成其他模型只需改这一行
)

# 非流式调用
response = client.chat.completions.create(
    model="deepseek-chat",
    messages=[
        {"role": "system", "content": "你是一位糖尿病健康顾问"},
        {"role": "user", "content": "我刚吃了两碗热干面，血糖 16.8"}
    ]
)
print(response.choices[0].message.content)

# 流式调用（SSE）
stream = client.chat.completions.create(
    model="deepseek-chat",
    messages=[...],
    stream=True
)
for chunk in stream:
    if chunk.choices[0].delta.content:
        print(chunk.choices[0].delta.content, end="", flush=True)
""", "代码示例：")

    h2(doc, "7.2 System Prompt 工程")
    para(doc, "System Prompt 是控制 LLM 行为的核心。好的 Prompt 需要：")
    bullet(doc, "角色定义：明确 AI 是谁（'你是 SugarClaw 代谢教练'）")
    bullet(doc, "知识注入：将指南/用户档案/食物数据嵌入 prompt")
    bullet(doc, "输出约束：限制回复格式、长度、措辞（'不使用绝对禁止等限制性指令'）")
    bullet(doc, "动态上下文：根据用户问题检索相关指南条目追加到 prompt")

    h2(doc, "7.3 多 Agent 架构")
    para(doc, "SugarClaw 定义了三个 Agent 角色，通过加权优先级协同：")
    bullet(doc, "生理罗盘（权重 1.0）：血糖数据分析、卡尔曼预测、预警生成")
    bullet(doc, "地道风味（权重 0.8）：食物 GI/GL 查询、对冲方案推荐")
    bullet(doc, "心理防线（权重 0.6）：情绪识别、认知重构、避免限制性语言")
    para(doc, "当建议冲突时（如预测低血糖时不应推荐运动），高权重 Agent 的判断优先。")

    resource_box(doc, [
        "OpenAI API 文档: https://platform.openai.com/docs/",
        "DeepSeek API 文档: https://platform.deepseek.com/api-docs/",
        "Prompt Engineering Guide: https://www.promptingguide.ai/zh",
    ])
    doc.add_page_break()

    # ═══════════════════════════════════════
    # 第八章
    # ═══════════════════════════════════════
    h1(doc, "第八章  PubMed 文献检索集成")

    h2(doc, "8.1 NCBI E-Utilities 三大接口")
    para(doc, "PubMed 是全球最大的生物医学文献数据库。NCBI 提供 E-Utilities API 进行程序化检索。")
    bullet(doc, "esearch：搜索关键词 → 返回文章 ID 列表")
    bullet(doc, "esummary：根据 ID 获取文章摘要信息（标题、作者、期刊、日期）")
    bullet(doc, "efetch：获取完整摘要文本（Abstract）")

    h2(doc, "8.2 实现示例")
    code_block(doc, """
import urllib.request, urllib.parse, xml.etree.ElementTree as ET

BASE = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils"

def search_pubmed(query, max_results=5):
    # Step 1: esearch — 获取文章 ID
    params = urllib.parse.urlencode({
        "db": "pubmed", "term": query,
        "retmax": max_results, "sort": "relevance"
    })
    url = f"{BASE}/esearch.fcgi?{params}"
    xml = urllib.request.urlopen(url).read()
    root = ET.fromstring(xml)
    ids = [id_elem.text for id_elem in root.findall(".//Id")]

    # Step 2: esummary — 获取摘要信息
    if not ids: return []
    params2 = urllib.parse.urlencode({"db": "pubmed", "id": ",".join(ids)})
    xml2 = urllib.request.urlopen(f"{BASE}/esummary.fcgi?{params2}").read()
    root2 = ET.fromstring(xml2)

    articles = []
    for doc in root2.findall(".//DocSum"):
        pmid = doc.find("Id").text
        title = ""
        for item in doc.findall("Item"):
            if item.get("Name") == "Title":
                title = item.text
        articles.append({"pmid": pmid, "title": title})
    return articles

# 使用
results = search_pubmed("diabetes CGM kalman filter")
for a in results:
    print(f"[{a['pmid']}] {a['title']}")
""", "完整可运行代码：")

    h2(doc, "8.3 速率限制")
    para(doc, "NCBI 限制：无 API Key 时 3 请求/秒，有 Key 时 10 请求/秒。"
         "SugarClaw 实现了指数退避重试（429/5xx 时等待后重试）。")

    resource_box(doc, [
        "NCBI E-Utilities 文档: https://www.ncbi.nlm.nih.gov/books/NBK25501/",
        "PubMed API Key 申请: https://www.ncbi.nlm.nih.gov/account/",
    ])
    doc.add_page_break()

    # ═══════════════════════════════════════
    # 第九章
    # ═══════════════════════════════════════
    h1(doc, "第九章  BLE 蓝牙协议与 CGM 数据")

    h2(doc, "9.1 BLE 基础概念")
    para(doc, "BLE（Bluetooth Low Energy）是低功耗蓝牙协议，CGM 设备通过 BLE 将血糖数据传输到手机。"
         "核心概念：")
    bullet(doc, "GATT（通用属性配置文件）：定义数据的组织方式 — Service → Characteristic → Descriptor")
    bullet(doc, "Glucose Service UUID: 0x1808 — Bluetooth SIG 标准定义的血糖服务")
    bullet(doc, "Measurement Characteristic: 0x2AA7 — 包含实际血糖读数的特征值")

    h2(doc, "9.2 IEEE 11073 SFLOAT 数据格式")
    para(doc, "CGM 传感器传输的血糖值使用 SFLOAT（Short Float）编码：2 字节 = 4 位指数 + 12 位尾数。")
    code_block(doc, """
def parse_sfloat(raw_bytes):
    \"\"\"解析 IEEE 11073 SFLOAT (16-bit)\"\"\"
    value = int.from_bytes(raw_bytes, byteorder='little')
    # 高 4 位 = 指数（有符号）
    exponent = (value >> 12) & 0x0F
    if exponent >= 8: exponent -= 16  # 补码转有符号
    # 低 12 位 = 尾数（有符号）
    mantissa = value & 0x0FFF
    if mantissa >= 2048: mantissa -= 4096
    # 特殊值检测
    if mantissa == 0x07FF: return float('nan')
    if mantissa == 0x0800: return float('inf')
    return mantissa * (10.0 ** exponent)

# 示例: 0x00B4 → exponent=0, mantissa=180 → 180 mg/dL → 10.0 mmol/L
print(parse_sfloat(b'\\xb4\\x00'))  # 180.0
""", "SFLOAT 解析代码：")

    h2(doc, "9.3 24h 模拟数据生成")
    para(doc, "SugarClaw 的 generate_demo_data() 生成逼真的 24 小时 CGM 数据（288 个点，每 5 分钟）：")
    bullet(doc, "基线血糖 ~6.0 mmol/L")
    bullet(doc, "黎明现象（4:00-6:30）：皮质醇升高导致血糖 +1.5 mmol/L")
    bullet(doc, "三餐峰值：早餐 GI=75、午餐 GI=60、晚餐 GI=70")
    bullet(doc, "餐后胰岛素作用：4U 速效胰岛素，ISF=0.73，tau=77min 指数衰减")
    bullet(doc, "高斯噪声 sigma=0.3 mmol/L，生理钳制 2.2-22.2 mmol/L")

    resource_box(doc, [
        "Bluetooth GATT Specifications: https://www.bluetooth.com/specifications/specs/",
        "IEEE 11073-20601: Health Informatics Standard",
    ])
    doc.add_page_break()

    # ═══════════════════════════════════════
    # 第十章
    # ═══════════════════════════════════════
    h1(doc, "第十章  部署与运维")

    h2(doc, "10.1 本地开发环境搭建")
    code_block(doc, """
# 1. Python 环境
python3 -m venv venv
source venv/bin/activate
pip install fastapi uvicorn python-docx chromadb openai

# 2. Flutter SDK (https://docs.flutter.dev/get-started/install)
# macOS: 下载 Flutter SDK → 解压 → 添加到 PATH
export PATH="$PATH:$HOME/development/flutter/bin"
flutter doctor  # 检查环境

# 3. 启动后端
cd backend && python3 api.py  # 默认 8080 端口

# 4. 构建前端
cd frontend && flutter build web --release

# 5. 访问 http://localhost:8080
""", "步骤：")

    h2(doc, "10.2 Cloudflare Tunnel 内网穿透")
    para(doc, "让外网用户无需公网 IP 即可访问你的本地服务。原理：本地运行 cloudflared 客户端，"
         "通过 Cloudflare 的全球边缘网络建立加密隧道，分配一个公网域名。")
    code_block(doc, """
# 安装
brew install cloudflared

# 一键启动（免费，临时域名）
cloudflared tunnel --url http://localhost:8080
# 输出: https://xxxx-xxxx.trycloudflare.com

# 把这个链接发给任何人，他们就能访问你的 SugarClaw！
""", "使用方法：")

    h2(doc, "10.3 进程管理")
    bullet(doc, "开发环境：直接 python3 api.py（终端关闭则停止）")
    bullet(doc, "后台运行：nohup python3 api.py > /tmp/sugarclaw.log 2>&1 &")
    bullet(doc, "生产环境：使用 systemd 或 supervisor 管理进程自动重启")
    doc.add_page_break()

    # ═══════════════════════════════════════
    # 附录 A
    # ═══════════════════════════════════════
    h1(doc, "附录 A  推荐学习路径（12 周规划）")

    weeks = [
        ("Week 1-2", "Python 基础 + FastAPI 入门",
         "学习 Python 语法、函数、类、异步；搭建第一个 FastAPI 项目，实现 CRUD API"),
        ("Week 3-4", "SQLite + 完整 REST API",
         "学习 SQL 基础、sqlite3 模块；为 API 添加数据库持久化；学习 Pydantic 数据验证"),
        ("Week 5-6", "Flutter / Dart 入门 + Provider",
         "学习 Dart 语法；掌握 Flutter Widget 树、布局系统；用 Provider 实现状态管理"),
        ("Week 7-8", "卡尔曼滤波理论 + 实现",
         "理解贝叶斯估计直觉；推导 KF 5 个方程；用 numpy 实现 KF/EKF；应用到血糖数据"),
        ("Week 9-10", "ChromaDB + RAG 系统",
         "理解向量嵌入；用 ChromaDB 构建食物搜索；实现 RAG 架构连接 LLM"),
        ("Week 11-12", "LLM 集成 + 部署上线",
         "调用 DeepSeek API；设计 System Prompt；实现 SSE 流式 Chat；用 Cloudflare Tunnel 部署"),
    ]
    for week, topic, detail in weeks:
        h2(doc, f"{week}：{topic}")
        para(doc, detail)

    doc.add_page_break()

    # ═══════════════════════════════════════
    # 附录 B
    # ═══════════════════════════════════════
    h1(doc, "附录 B  参考资源汇总")

    h2(doc, "书籍")
    bullet(doc, "《Kalman and Bayesian Filters in Python》— Roger Labbe（免费在线）")
    bullet(doc, "《Flask Web 开发》— Miguel Grinberg（FastAPI 理念相通）")
    bullet(doc, "《Flutter 实战》— 杜文（中文 Flutter 入门经典）")

    h2(doc, "在线课程")
    bullet(doc, "freeCodeCamp: FastAPI Course (YouTube, 免费)")
    bullet(doc, "Flutter 官方 Codelabs (免费)")
    bullet(doc, "Coursera: Machine Learning Specialization — Andrew Ng")

    h2(doc, "文档")
    bullet(doc, "FastAPI: https://fastapi.tiangolo.com/zh/")
    bullet(doc, "Flutter: https://docs.flutter.dev/")
    bullet(doc, "ChromaDB: https://docs.trychroma.com/")
    bullet(doc, "SQLite: https://www.sqlite.org/docs.html")
    bullet(doc, "DeepSeek API: https://platform.deepseek.com/api-docs/")
    bullet(doc, "PubMed E-Utilities: https://www.ncbi.nlm.nih.gov/books/NBK25501/")

    h2(doc, "论文")
    bullet(doc, "Clarke WL et al. 'Evaluating Clinical Accuracy of Systems for Self-Monitoring of Blood Glucose' Diabetes Care, 1987")
    bullet(doc, "Turksoy K et al. 'Multivariable Adaptive Identification and Control for Artificial Pancreas Systems' IEEE Trans Biomed Eng, 2014")
    bullet(doc, "Oviedo S et al. 'A Review of Personalized Blood Glucose Prediction Strategies for T1DM' Int J Numer Method Biomed Eng, 2017")

    # 保存
    doc.save(OUTPUT)
    print(f"技术学习指南已保存至: {OUTPUT}")


if __name__ == "__main__":
    build()
