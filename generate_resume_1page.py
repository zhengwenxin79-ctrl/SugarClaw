#!/usr/bin/env python3
"""Generate a ONE-PAGE resume as a Word document."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Style constants ──
BLUE = RGBColor(0x2E, 0x86, 0xC1)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x55, 0x55, 0x55)
FONT_CN = "微软雅黑"
FONT_EN = "Calibri"

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)


def set_font(run, size, color=BLACK, bold=False, cn=FONT_CN, en=FONT_EN):
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.font.name = en
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), cn)


def set_spacing(paragraph, before=0, after=0, line=None):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = Pt(line)


def add_tab_stop_right(paragraph, pos_cm=15.0):
    """Add a right-aligned tab stop."""
    pPr = paragraph._element.get_or_add_pPr()
    tabs = pPr.find(qn("w:tabs"))
    if tabs is None:
        tabs = parse_xml(f'<w:tabs {nsdecls("w")} />')
        pPr.append(tabs)
    tab = parse_xml(f'<w:tab {nsdecls("w")} w:val="right" w:pos="{int(pos_cm * 567)}" />')
    tabs.append(tab)


# ── Header: Name ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(p, before=0, after=0, line=22)
r = p.add_run("雯馨 郑")
set_font(r, 18, BLACK, bold=True)

# ── Subtitle ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(p, before=0, after=0, line=12)
r = p.add_run("生物医学硕士研究生 · AI+ 医学交叉方向 | 吉林大学，长春")
set_font(r, 9, BLUE)

# ── Contact ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(p, before=0, after=2, line=12)
r = p.add_run("(+86) 178-0805-8287 | zhengwenxin79@gmail.com | zhengwenxin")
set_font(r, 8.5, GRAY)


def add_section_heading(text):
    p = doc.add_paragraph()
    set_spacing(p, before=4, after=1, line=14)
    r = p.add_run(text)
    set_font(r, 12, BLUE, bold=True)
    # Bottom border
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="2E86C1"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def add_entry_title(left_text, right_text=None):
    p = doc.add_paragraph()
    set_spacing(p, before=2, after=0, line=13)
    if right_text:
        add_tab_stop_right(p)
    r = p.add_run(left_text)
    set_font(r, 10, BLACK, bold=True)
    if right_text:
        r2 = p.add_run(f"\t{right_text}")
        set_font(r2, 8.5, GRAY)
    return p


def add_bullet(text, indent_cm=0.5):
    p = doc.add_paragraph()
    set_spacing(p, before=1, after=0, line=12)
    pf = p.paragraph_format
    pf.left_indent = Cm(indent_cm)
    pf.first_line_indent = Cm(-0.3)
    r = p.add_run("• ")
    set_font(r, 8.5, BLACK)
    r2 = p.add_run(text)
    set_font(r2, 8.5, BLACK)
    return p


def add_plain_line(text, size=8.5, color=BLACK, indent_cm=0.0, before=1, after=0, line=12):
    p = doc.add_paragraph()
    set_spacing(p, before=before, after=after, line=line)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    r = p.add_run(text)
    set_font(r, size, color)
    return p


# ══════════════════════════════════════
# 教育背景
# ══════════════════════════════════════
add_section_heading("教育背景")

p = add_entry_title("吉林大学 · 基础医学院")
# add date on same line
add_tab_stop_right(p)
r = p.add_run("\t长春 | 2024.09–至今")
set_font(r, 8.5, GRAY)

# sub-info
p2 = doc.add_paragraph()
set_spacing(p2, before=0, after=0, line=12)
p2.paragraph_format.left_indent = Cm(0.2)
r = p2.add_run("细胞生物学硕士（在读）| 研究方向：抗菌材料在糖尿病创面修复中的应用；参与AI微专业计划；SCI论文1篇（在投）")
set_font(r, 8.5, BLACK)

# ══════════════════════════════════════
# 专业技能
# ══════════════════════════════════════
add_section_heading("专业技能")

skills = [
    "生物医学：细胞培养、分子生物学实验、抗菌材料表征、糖尿病动物模型",
    "AI与全栈开发：Python、PyTorch、RAG、卡尔曼滤波 | FastAPI、Flutter、SQLite、ChromaDB、SSE实时流",
    "工具与语言：Git、Linux、LaTeX、OpenClaw、Docker | 英语 CET-6",
]
for s in skills:
    add_bullet(s, indent_cm=0.5)

# ══════════════════════════════════════
# 项目与研究经历
# ══════════════════════════════════════
add_section_heading("项目与研究经历")

# ── Project 1: SugarClaw ──
p = add_entry_title("SugarClaw — AI驱动的糖尿病代谢决策引擎")
add_tab_stop_right(p)
r = p.add_run("\t长春 | 独立开发者 | 2026.02–至今")
set_font(r, 8.5, GRAY)

sc_bullets = [
    "独立开发全栈糖尿病智能管理平台（FastAPI + Flutter），7,700行代码，支持Web/iOS/Android多端部署",
    "实现KF/EKF/UKF三变体自适应卡尔曼滤波器，基于125名患者128K条CGM数据校准，Clarke A准确率95.68%",
    "构建501种中国食物GI/GL向量数据库（ChromaDB语义检索，36地域×18类别），支持方言/俗名智能识别",
    "设计「对冲天平」交互范式（风险量化+配菜/运动对冲），集成CDS/ADA/IDF/EASD/WHO五大权威指南知识库",
    "构建多Agent协作架构（生理罗盘·地道风味·心理防线），集成PubMed实时文献检索与BLE CGM协议解析",
]
for b in sc_bullets:
    add_bullet(b)

# ── Project 2: Nano-medicine ──
p = add_entry_title("纳米医学课题组 · 吉林大学基础医学院")
add_tab_stop_right(p)
r = p.add_run("\t长春 | 硕士研究生 | 2024.09–至今")
set_font(r, 8.5, GRAY)

nano_bullets = [
    "研究CuSe/CuFeSe\u2082复合抗菌材料的抗菌性能，建立小鼠皮肤细菌感染模型评估耐药菌杀菌活性",
    "参与Janus结构贵金属@CeO\u2082纳米酶课题，论文投稿至Biomaterials（在投）",
]
for b in nano_bullets:
    add_bullet(b)

# ── Project 3: Bio-mimetic ──
p = add_entry_title("仿生医疗器械实验室 · 南华大学")
add_tab_stop_right(p)
r = p.add_run("\t衡阳 | 本科研究助理 | 2023.01–2023.06")
set_font(r, 8.5, GRAY)

add_bullet("设计仿生蛇牙微针阵列用于糖尿病无痛胰岛素透皮给药，采用光固化树脂优化针体结构")

# ══════════════════════════════════════
# 学生工作
# ══════════════════════════════════════
add_section_heading("学生工作")
add_plain_line(
    "吉林大学研究生行政助理（2024.09–2025.09）| 南华大学新媒体部副部长（2020.09–2022.06，公众号阅读量50万+）| 南华大学女排队长（新生杯冠军）",
    size=8.5, color=BLACK, indent_cm=0.2, before=1, after=0, line=12,
)

# ══════════════════════════════════════
# 荣誉奖项
# ══════════════════════════════════════
add_section_heading("荣誉奖项")
add_plain_line(
    "AI微专业优秀结业项目（2025）| 国家励志奖学金（2022）| 校级优秀学生干部（2021）",
    size=8.5, color=BLACK, indent_cm=0.2, before=1, after=0, line=12,
)

# ── Save ──
output_path = "/Users/zwx/Downloads/郑雯馨_简历_单页版.docx"
doc.save(output_path)
print(f"Resume saved to {output_path}")
