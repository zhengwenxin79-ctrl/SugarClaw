"""
根据原简历格式，将 SugarClaw 项目写入更新后的简历。
输出为 Word 文档，尽量还原原 PDF 的排版风格。
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "郑雯馨_简历_SugarClaw.docx")

# ─── 样式常量 ───────────────────────────────
BLUE = RGBColor(0x2E, 0x86, 0xC1)    # 标题蓝
BLACK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x55, 0x55, 0x55)
FONT_CN = "微软雅黑"
FONT_EN = "Calibri"


def set_font(run, size=10.5, bold=False, color=BLACK, cn=FONT_CN, en=FONT_EN):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = en
    run._element.rPr.rFonts.set(qn("w:eastAsia"), cn)


def add_heading_with_line(doc, text, level=1):
    """添加带下划线的蓝色大标题"""
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    p.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=16, bold=True, color=BLUE)
    # 底部边框线
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "8",
        qn("w:space"): "1",
        qn("w:color"): "2E86C1",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_entry_header(doc, title, location, subtitle="", date=""):
    """项目/经历标题行：左侧标题+右侧地点，第二行左侧角色+右侧时间"""
    # 第一行
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(0)
    run = p.add_run(title)
    set_font(run, size=11, bold=True, color=BLACK)
    if location:
        run2 = p.add_run(f"\t{location}")
        set_font(run2, size=10, bold=False, color=GRAY)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(16.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    # 第二行（副标题 + 日期）
    if subtitle or date:
        p2 = doc.add_paragraph()
        p2.space_before = Pt(0)
        p2.space_after = Pt(2)
        if subtitle:
            run3 = p2.add_run(subtitle)
            set_font(run3, size=9.5, bold=False, color=GRAY)
        if date:
            run4 = p2.add_run(f"\t{date}")
            set_font(run4, size=9.5, bold=False, color=GRAY)
            p2.paragraph_format.tab_stops.add_tab_stop(Cm(16.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    return p


def add_bullet(doc, text, indent_cm=0.8):
    """添加一条项目符号"""
    p = doc.add_paragraph()
    p.space_before = Pt(0)
    p.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(indent_cm)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    run = p.add_run("• " + text)
    set_font(run, size=10, bold=False, color=BLACK)
    return p


def add_bullet_bold_prefix(doc, prefix, text, indent_cm=0.8):
    """带加粗前缀的项目符号"""
    p = doc.add_paragraph()
    p.space_before = Pt(0)
    p.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(indent_cm)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    run1 = p.add_run("• " + prefix)
    set_font(run1, size=10, bold=True, color=BLACK)
    run2 = p.add_run(text)
    set_font(run2, size=10, bold=False, color=BLACK)
    return p


def build():
    doc = Document()

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # ═══════════════════════════════════════
    # 头部信息
    # ═══════════════════════════════════════
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(2)
    run = p.add_run("雯馨 郑")
    set_font(run, size=22, bold=True, color=BLACK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(0)
    p.space_after = Pt(2)
    run = p.add_run("生物医学硕士研究生  ·  AI+ 医学交叉方向")
    set_font(run, size=11, bold=False, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(0)
    p.space_after = Pt(2)
    run = p.add_run("吉林大学，长春")
    set_font(run, size=10, color=GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(0)
    p.space_after = Pt(2)
    run = p.add_run("(+86) 178-0805-8287  |  zhengwenxin79@gmail.com  |  zhengwenxin")
    set_font(run, size=9.5, color=GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(2)
    p.space_after = Pt(6)
    run = p.add_run('"优雅从容，向阳而生。"')
    set_font(run, size=10, bold=False, color=GRAY)

    # ═══════════════════════════════════════
    # 教育背景
    # ═══════════════════════════════════════
    add_heading_with_line(doc, "教育背景")

    add_entry_header(doc, "吉林大学·基础医学院", "长春",
                     "细胞生物学硕士（在读）", "2024.09 - 至今")
    add_bullet(doc, "研究方向：抗菌材料在糖尿病创面修复中的应用")
    add_bullet(doc, "参与 AI 微专业计划，系统学习机器学习、深度学习及其生物医学应用")
    add_bullet(doc, "在读期间发表 SCI 论文 1 篇（在投）")

    add_entry_header(doc, "南华大学·电气工程学院", "衡阳",
                     "生物医学工程工学学士", "2019.09 - 2024.06")
    add_bullet(doc, "GPA: 3.7/4.0，专业排名前 15%")
    add_bullet(doc, "核心课程：信号与系统、医学仪器原理、医学成像技术、生物医学传感器、数字信号处理")
    add_bullet(doc, "获国家励志奖学金、校级优秀学生干部")

    # ═══════════════════════════════════════
    # 专业技能
    # ═══════════════════════════════════════
    add_heading_with_line(doc, "专业技能")

    skills = [
        ("生物医学", "细胞培养、分子生物学实验、抗菌材料表征、糖尿病动物模型、慢性创面处理"),
        ("数据科学与 AI", "Python、机器学习（sklearn）、深度学习（PyTorch）、RAG 技术、卡尔曼滤波、医学数据分析"),
        ("全栈开发", "FastAPI、Flutter（Web/Mobile）、SQLite、ChromaDB 向量检索、SSE 实时流、RESTful API 设计"),
        ("工具与平台", "LaTeX、Git、Linux、OpenClaw、Docker、Cloudflare Tunnel 部署"),
        ("语言能力", "中文（母语）、英语（CET-6，学术论文读写流利）"),
    ]
    for label, content in skills:
        add_bullet_bold_prefix(doc, f"{label}  ", content)

    # ═══════════════════════════════════════
    # 项目与研究经历
    # ═══════════════════════════════════════
    add_heading_with_line(doc, "项目与研究经历")

    # ── SugarClaw（核心项目，重点展开）──
    add_entry_header(doc, "SugarClaw — AI 驱动的糖尿病代谢决策引擎", "长春",
                     "独立开发者 · 全栈设计与实现", "2026.02 - 至今")

    add_bullet(doc,
        "独立设计并开发全栈糖尿病智能管理平台（Python 后端 + Flutter 前端），"
        "总代码量约 7,700 行，支持 Web / iOS / Android 多端部署")
    add_bullet(doc,
        "实现三变体自适应卡尔曼滤波器（KF/EKF/UKF），根据血糖波动自动切换滤波模式；"
        "基于 125 名患者、128,157 条 CGM 数据进行参数校准，Clarke A 格点准确率达 95.68%（RMSE 12.35 mg/dL）")
    add_bullet(doc,
        "构建中国地域特色食物 GI/GL 向量数据库（501 种食物 × 36 地域 × 18 类别），"
        "基于 ChromaDB 实现语义检索，支持方言/俗名识别（如'过早吃碗面' → 热干面 GI=82）")
    add_bullet(doc,
        "设计「对冲天平」创新交互范式：左盘量化食物升糖风险（GI/GL/宏量营养素 + 时段修正），"
        "右盘智能推荐低 GI 配菜与运动对冲方案，物理倾斜角可视化平衡状态")
    add_bullet(doc,
        "集成 5 大权威临床指南知识库（CDS 2024 / ADA 2025 / IDF / EASD / WHO，共 34 条循证条目），"
        "Chat 对话中自动检索相关指南并标注来源与推荐等级")
    add_bullet(doc,
        "构建多 Agent 协作架构（生理罗盘·地道风味·心理防线），"
        "通过加权优先级调度避免矛盾建议，贯彻「认知重构而非限制」的干预理念")
    add_bullet(doc,
        "集成 PubMed E-Utilities 实时文献检索（4 种预设模式 + 自定义查询），"
        "支持摘要获取、历史缓存，为建议提供循证支撑")
    add_bullet(doc,
        "实现 BLE CGM 蓝牙协议解析（GATT 0x1808 / IEEE 11073 SFLOAT），"
        "支持 24h 模拟数据生成与 SSE 实时推流可视化")

    # ── 纳米医学课题 ──
    add_entry_header(doc, "纳米医学课题组·吉林大学基础医学院", "长春",
                     "硕士研究生", "2024.09 - 至今")
    add_bullet(doc, "参与「Janus 结构贵金属 @CeO₂ 纳米酶」相关课题研究")
    add_bullet(doc, "研究 CuSe/CuFeSe2 复合抗菌材料的体内外抗菌性能")
    add_bullet(doc, "建立小鼠皮肤细菌感染模型，评估材料对耐药菌的杀菌活性")
    add_bullet(doc, "撰写论文投稿至 Bacterial Materials（在投）")

    # ── 仿生医疗器械 ──
    add_entry_header(doc, "仿生医疗器械实验室·南华大学", "衡阳",
                     "本科研究助理", "2023.01 - 2023.06")
    add_bullet(doc, "设计并制备仿生蛇牙微针阵列，用于糖尿病患者无痛胰岛素透皮给药")
    add_bullet(doc, "采用光固化树脂材料优化针体结构，实现快速皮肤穿刺与微创释药")
    add_bullet(doc, "通过力学测试验证穿刺可靠性，提升患者长期胰岛素治疗依从性")

    # ── 医学影像 ──
    add_entry_header(doc, "医学影像智能分析系统", "长春",
                     "AI 课程项目", "2025.03 - 2025.06")
    add_bullet(doc, "基于 ResNet 构建糖尿病视网膜病变（DR）分级模型，准确率达 92%")
    add_bullet(doc, "使用 Grad-CAM 可视化病变区域，辅助临床医生定位微血管瘤")
    add_bullet(doc, "项目获校级 AI 微专业优秀结业项目")

    # ═══════════════════════════════════════
    # 学生工作与社会活动
    # ═══════════════════════════════════════
    add_heading_with_line(doc, "学生工作与社会活动")

    add_entry_header(doc, "吉林大学基础医学院", "长春",
                     "研究生行政助理", "2024.09 - 2025.09")
    add_bullet(doc, "协助学院日常行政事务，统筹安排学术会议与讲座")
    add_bullet(doc, "参与研究生招生宣传，接待来访学生与家长")

    add_entry_header(doc, "南华大学官方微信公众号", "衡阳",
                     "新媒体部副部长", "2020.09 - 2022.06")
    add_bullet(doc, "负责校园新闻采编与公众号内容策划，阅读量累计超 50 万")
    add_bullet(doc, "统筹招生宣传专题，制作图文、短视频等多形式内容")

    add_entry_header(doc, "南华大学", "衡阳",
                     "女子排球队队长", "2019.09 - 2021.06")
    add_bullet(doc, "带领校队获得「新生杯」排球赛冠军")
    add_bullet(doc, "组织日常训练，培养团队协作与抗压能力")

    # ═══════════════════════════════════════
    # 荣誉奖项
    # ═══════════════════════════════════════
    add_heading_with_line(doc, "荣誉奖项")

    awards = [
        ("2025", "AI 微专业优秀结业项目", "吉林大学", "长春"),
        ("2022", "国家励志奖学金", "教育部", ""),
        ("2021", "校级优秀学生干部", "南华大学", "衡阳"),
        ("2019", "「新生杯」排球赛冠军", "南华大学", "衡阳"),
    ]
    for year, title, org, loc in awards:
        p = doc.add_paragraph()
        p.space_before = Pt(1)
        p.space_after = Pt(1)
        run1 = p.add_run(f"{year}  ")
        set_font(run1, size=10, bold=False, color=GRAY)
        run2 = p.add_run(title)
        set_font(run2, size=10, bold=True, color=BLACK)
        run3 = p.add_run(f", {org}")
        set_font(run3, size=10, bold=False, color=BLACK)
        if loc:
            run4 = p.add_run(f"\t{loc}")
            set_font(run4, size=10, bold=False, color=GRAY)
            p.paragraph_format.tab_stops.add_tab_stop(Cm(16.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    # ═══════════════════════════════════════
    doc.save(OUTPUT)
    print(f"简历已保存至: {OUTPUT}")


if __name__ == "__main__":
    build()
